# -*- coding: utf-8 -*-
"""
通用智能会议纪要生成工具 - 主程序（接入大模型）
适配所有会议类型，真实调用大模型实现智能提取，输出飞书格式纪要
"""
import streamlit as st
import os
from modules.preprocess import parse_speech
from modules.extract import extract_meeting_info
from modules.template import fill_template, load_all_templates
from modules.output import save_md, save_word

# ====================== 页面基础配置 =======================
st.set_page_config(
    page_title="通用智能会议纪要生成工具",
    page_icon="📝",
    layout="wide"
)
st.title("📝 通用智能会议纪要生成工具")
st.subheader("上传任意文字会议记录，生成飞书风格智能纪要", divider="blue")

# ====================== 侧边栏：API配置+模板选择 =======================
st.sidebar.title("⚙️ 工具配置")
# 1. 通义千问API Key（必填，否则无法智能提取）
api_key = st.sidebar.text_input("通义千问API Key（必填）", type="password", 
                                help="从阿里云百炼控制台获取：https://dashscope.aliyun.com/")
# 2. 会议模板选择
st.sidebar.subheader("📋 选择会议模板")
template_type = st.sidebar.selectbox(
    "适配所有会议类型",
    options=["通用商务会议", "项目同步会议", "需求评审会议", "周度例会"],
    index=0
)
# 加载选中的模板
templates = load_all_templates()
selected_template = templates[template_type]

# ====================== 主界面：文件上传+智能提取 ======================
# 支持TXT文件上传（真实读取文件内容）
uploaded_file = st.file_uploader("上传会议文字记录（TXT格式）", type=["txt"])
if uploaded_file is not None:
    # 真实读取上传的文件内容
    try:
        meeting_text = uploaded_file.read().decode("utf-8")
        st.success("✅ 文件上传成功！")
        
        # 预览原始文本
        with st.expander("📄 查看原始会议记录", expanded=False):
            st.text(meeting_text)
        
        # 生成纪要按钮
        if st.button("🚀 生成飞书风格智能纪要", type="primary"):
            if not api_key:
                st.error("❌ 请先在侧边栏输入通义千问API Key！")
            else:
                with st.spinner("🔍 正在调用大模型分析会议内容..."):
                    # 1. 预处理文本
                    speech_list = parse_speech(meeting_text)
                    # 2. 智能提取（传递API Key）
                    extract_result = extract_meeting_info(speech_list, template_type, api_key)
                    
                    # 3. 处理提取结果
                    if "error" in extract_result:
                        st.error(f"❌ {extract_result['error']}")
                    else:
                        # 4. 填充飞书模板
                        summary_text = fill_template(extract_result, selected_template)
                        # 5. 预览纪要
                        st.subheader("📋 飞书风格智能会议纪要", divider="green")
                        st.markdown(summary_text)
                        
                        # 6. 生成下载文件
                        md_path = save_md(summary_text, f"{template_type}_飞书风格纪要.md")
                        word_path = save_word(summary_text, f"{template_type}_飞书风格纪要.docx")
                        
                        # 7. 下载按钮（真实读取文件）
                        with open(md_path, 'w', encoding='utf-8') as f:
                            f.write(summary_text)
                        with open(md_path, 'r', encoding='utf-8') as f:
                            st.download_button(f"📥 下载MD格式-{template_type}纪要", f, file_name=md_path)
                        
                        # Word文件生成（真实写入）
                        from docx import Document
                        doc = Document()
                        doc.add_heading(extract_result["会议核心信息"]["主题"], level=1)
                        for line in summary_text.split("\n"):
                            if line.startswith("##"):
                                doc.add_heading(line.replace("## ", ""), level=2)
                            elif line.startswith("-"):
                                doc.add_paragraph(line, style='List Bullet')
                            elif line:
                                doc.add_paragraph(line)
                        doc.save(word_path)
                        with open(word_path, 'rb') as f:
                            st.download_button(f"📥 下载Word格式-{template_type}纪要", f, file_name=word_path)
                        
                        st.success(f"🎉 {template_type}纪要生成完成！完全匹配飞书格式")
    except Exception as e:
        st.error(f"❌ 文件读取失败：{str(e)}")

# ====================== 页脚（兼容所有Streamlit版本）======================
st.markdown(
    """
    <div style='text-align: center; color: #666; font-size: 12px; margin-top: 50px;'>
    💡 工具基于通义千问大模型开发 | 输出飞书风格标准化纪要 | 支持所有办公会议类型
    </div>
    """,
    unsafe_allow_html=True
)
